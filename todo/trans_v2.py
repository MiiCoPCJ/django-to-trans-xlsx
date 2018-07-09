import os, sys
import re
from docx import *
from win32com.client import Dispatch, constants, gencache
from win32com import client



# 执行文件的路径
path = os.path.split(os.path.realpath(__file__))[0]


def make_docx(pre,file_path,document):
    str = ''

    f = open(file_path,'r',encoding='utf-8')
    for line in f.readlines():
        line = line.strip()
        # 去除#注明部分
        if not len(line) or line.startswith('#'):
            continue
        str += line + '\n'
    f.close()

    # 去除"""注释
    str = re.sub('"""(.|\n)*?"""','',str)

    # 正则 查找
    pattern = re.compile(r'class\s+\w+\((AbstractBaseUser|models.Model|MPTTModel)\):.*?verbose_name\s?=\s?verbose_name_plural\s?=\s?(\'|")(\w|-)+(\'|")', re.S)

    find = re.finditer(pattern,str)

    # 历遍
    for m in find:
        model = m.group()

        # 表名 英文
        data_name = re.search('class\s+(\w+)\(.*?\):', model).group(1).lower()
        # 表名 中文
        data_name_cn = re.search('verbose_name_plural\s?=\s?(\'|")((\w|-)+)(\'|")', model).group(2)

        p = document.add_paragraph(data_name+'('+data_name_cn+')',style='Heading 1')

        # 选择字段说明
        dic = {}
        i = 0
        pattern = re.compile(r'\w+\s?=\s?\(.*?\)\n',re.S)
        chois = re.finditer(pattern,model)
        for cho in chois:

            b = cho.group()
            sub = {}
            lts = re.finditer('\(\s?(\w+)\s?,\s?(\'|")(.*?)(\'|")\s?\)',b)
            for lt in lts:

                choice_name = lt.group(1)
                choice_name_cn = lt.group(3)

                choice_no = re.search(choice_name+'\s?=\s?(\d)',model).group(1)

                sub.update({choice_name:[choice_name_cn,choice_no]})

            dic.update({i:sub})
            i = i + 1


        # 获取字段部分
        field = re.search('class\s+\w+\(.*?\):(.*)class\sMeta:', model, re.S).group(1)
        stp = field.strip().split('\n')

        # 计算字段个数
        # k = len(re.findall('models\.',stp))

        # 添加表格
        table = document.add_table(rows=1, cols=5, style='Table Grid')

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '字段'
        hdr_cells[1].text = '类型'
        hdr_cells[2].text = '长度'
        hdr_cells[3].text = '默认值'
        hdr_cells[4].text = '描述'


        for sp in stp:
            # 外键处理
            key = re.search('(\w+)\s?=\s?models\.ForeignKey\((.*)\)',sp)
            if key is not None and len(key.group()):
                continue



            d = re.search('(\w+)\s?=\s?models\.(\w+)Field\((.*)\)',sp)

            if d is not None and len(d.group()):
                row_cells = table.add_row().cells

                field_name = d.group(1)
                field_type = d.group(2)


                if re.search('verbose_name\s?=\s?(\'|")((\w|-|\/)+)(\'|")', d.group(3)):
                    field_name_cn = re.search('verbose_name\s?=\s?(\'|")((\w|-|\/)+)(\'|")', d.group(3)).group(2)
                else:
                    field_name_cn = re.sub('("|\')','',d.group(3).split(',')[0])

                row_cells[0].text = field_name
                if field_type == 'Char':
                    row_cells[1].text = 'varchar'
                else:
                    row_cells[1].text = field_type.lower()


                if field_type == 'Decimal':
                    max_digits = re.search('max_digits\s?=\s?(\d+)\s?(,)?', d.group(3)).group(1)
                    decimal_places = re.search('decimal_places\s?=\s?(\d+)\s?(,)?', d.group(3)).group(1)
                    row_cells[2].text = max_digits+','+decimal_places
                elif field_type == 'Char':
                    max_length = re.search('max_length\s?=\s?(\d+)\s?(,)?', d.group(3)).group(1)
                    row_cells[2].text = max_length



                row_cells[4].text = row_cells[4].text + field_name_cn + '\n'

                # 默认值

                if re.search('default\s?=\s?(.*)(,)?',d.group(3)):
                    dft = re.search('default\s?=\s?(.*)(,)?',d.group(3)).group(1)
                    row_cells[3].text = dft
                    if dft:

                        f = -1
                        for i in dic:
                            if dft in dic[i]:
                                f = i
                                row_cells[3].text = dic[i][dft][1]

                        if f>=0:
                            for choice in dic[f]:
                                row_cells[4].text = row_cells[4].text + dic[f][choice][1]+' -- '+choice+'  '+dic[f][choice][0]+'\n'

                # 描述

                for dt in d.group(3).split(','):
                    de = re.search('(\w+)=.*',dt)
                    if de:
                        # if de.group(1) != 'verbose_name' and de.group(1)!='max_digits' and de.group(1)!='max_length' and de.group(1)!='default'\
                        #         and de.group(1) != 'decimal_places':
                        #     row_cells[4].text = row_cells[4].text + dt
                        if de.group(1) == 'auto_now':
                            row_cells[4].text = row_cells[4].text + '值:现今时间'
                        if de.group(1) == 'auto_now_add':
                            row_cells[4].text = row_cells[4].text + '值:现今时间'
                        if de.group(1) == 'choices':
                            row_cells[4].text = row_cells[4].text + dt


        for sp in stp:
            # 外键处理
            key = re.search('(\w+)\s?=\s?models\.ForeignKey\((.*)\)',sp)
            if key is not None and len(key.group()):
                row_cells = table.add_row().cells
                key_name = key.group(1)
                foreign = key.group(2).split(',')[0].strip()
                # 查找外键所在表
                link_key = re.search('from\s(\w+).models\simport\s' + foreign, str)
                if link_key is not None and len(link_key.group()):
                    dl = link_key.group(1) + '_' + foreign.lower()
                else:
                    dl = pre + '_' + foreign.lower()


                des = re.search('\s?' + foreign + '\s?,(.*)', key.group(2)).group(1).strip()
                row_cells[0].text = key_name + '_id'
                row_cells[4].text = '外键  \n外表('+dl+')  \n'


        document.add_paragraph('')


        pattern = re.compile(r'\w+\s?=\s?\(.*?\)\n', re.S)
        chois = re.finditer(pattern, model)
        for cho in chois:

            b = cho.group()
            stp = b.split('\n')
            for sp in stp:
                # 选择数据说明
                if re.findall(',', sp):
                    document.add_paragraph('     ' + sp, style='Caption')
                else:
                    document.add_paragraph(sp, style='Caption')


        document.add_paragraph('')
        document.add_paragraph('')









# 查找resource目录所有model文件
def loop_res():
    dir = path + "/../resource/"

    for parent, dirnames, filenames in os.walk(dir):
        for filename in filenames:
            file_path = os.path.join(parent, filename)

            pre = re.sub('_models.py','',filename)
            # 开启
            document = Document()
            make_docx(pre,file_path,document)
            document.save(path + '/../doc/word/'+pre+'.docx')

            # 转换为ptf
            doc2pdf(path + '/../doc/word/'+pre+'.docx', path + '/../doc/pdf/'+pre+'.pdf')


def doc2pdf(input, output):
    w = Dispatch('Word.Application')
    try:
        # 打开文件
        doc = w.Documents.Open(input, ReadOnly=1)
        # 转换文件
        doc.ExportAsFixedFormat(output, constants.wdExportFormatPDF,
                                Item=constants.wdExportDocumentWithMarkup, CreateBookmarks = constants.wdExportCreateHeadingBookmarks)
        return True
    except Exception as e:
        print(e)
        return False
    finally:
        w.Quit(constants.wdDoNotSaveChanges)




# Generate all the support we can.
def GenerateSupport():
    # enable python COM support for Word 2007
    # this is generated by: makepy.py -i "Microsoft Word 12.0 Object Library"
    gencache.EnsureModule('{00020905-0000-0000-C000-000000000046}', 0, 8, 4)


if __name__ == "__main__":
    loop_res()


    # document = Document()
    # make_docx('users', path+'/../resource/users_models.py', document)
    # document.save(path + '/../doc/model.docx')
    #
    # input = path + '/../doc/model.docx'
    # output = path + '/../doc/models.pdf'
    # GenerateSupport()
    # rc = doc2pdf(input, output)



    # document = Document()
    # styles = document.styles
    # for style in styles:
    #     print("'%s' -- %s" % (style.name, style.type))